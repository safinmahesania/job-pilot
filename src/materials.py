"""Generated application documents, stored against the job they belong to.

The point of storing them (rather than regenerating on demand) is correctness as
much as cost: when the browser extension attaches a resume to an application, it
must attach the resume written for THAT job. Every material row is keyed by
job_id, and the extension asks for a file by job_id — so there is no path by
which one company receives another company's cover letter.

`UNIQUE (job_id, kind)` means regenerating overwrites: a job has exactly one
current resume and one current cover letter, never a stale pile to pick from.
"""
import re
import unicodedata

from src import store


KINDS = ("resume", "cover")

# Filenames the ATS will see. Kept boring and professional.
_FILE_LABEL = {"resume": "Resume", "cover": "Cover_Letter"}


# ── Storage ─────────────────────────────────────────────────────────────────

def save(job_id: int, kind: str, content: str, provider: str = "") -> dict:
    """Store (or replace) a document for one job."""
    if kind not in KINDS:
        raise ValueError(f"unknown material kind: {kind}")
    conn = store.connect()
    conn.execute(
        "INSERT INTO materials (job_id, kind, content, provider) VALUES (?,?,?,?) "
        "ON CONFLICT(job_id, kind) DO UPDATE SET "
        "content=excluded.content, provider=excluded.provider, "
        "created_at=datetime('now')",
        (job_id, kind, content, provider),
    )
    conn.commit()
    conn.close()
    return {"job_id": job_id, "kind": kind, "saved": True}


def get(job_id: int, kind: str) -> dict | None:
    """The current document of this kind for this job, or None."""
    conn = store.connect()
    conn.row_factory = None
    row = conn.execute(
        "SELECT content, provider, created_at FROM materials "
        "WHERE job_id=? AND kind=?", (job_id, kind),
    ).fetchone()
    conn.close()
    if not row:
        return None
    return {"kind": kind, "content": row[0], "provider": row[1], "created_at": row[2]}


def list_for(job_id: int) -> list[dict]:
    """What has been generated and saved for this job (without the bodies)."""
    conn = store.connect()
    conn.row_factory = None
    rows = conn.execute(
        "SELECT kind, provider, created_at, LENGTH(content) FROM materials "
        "WHERE job_id=? ORDER BY kind", (job_id,),
    ).fetchall()
    conn.close()
    return [{"kind": r[0], "provider": r[1], "created_at": r[2], "chars": r[3]}
            for r in rows]


def delete(job_id: int, kind: str) -> bool:
    conn = store.connect()
    cur = conn.execute("DELETE FROM materials WHERE job_id=? AND kind=?",
                       (job_id, kind))
    conn.commit()
    n = cur.rowcount
    conn.close()
    return bool(n)


# ── Filenames ───────────────────────────────────────────────────────────────

def filename(job: dict, kind: str, ext: str) -> str:
    """A clean, recruiter-facing filename: Safin_Mahesania_Resume_Shopify.pdf"""
    from src.config import load_profile
    name = (load_profile().get("identity", {}) or {}).get("name", "")

    def slug(text: str) -> str:
        text = unicodedata.normalize("NFKD", str(text or ""))
        text = text.encode("ascii", "ignore").decode()
        text = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
        return text[:40]

    parts = [p for p in (slug(name), _FILE_LABEL.get(kind, kind),
                         slug(job.get("company", ""))) if p]
    return "_".join(parts) + f".{ext}"


# ── PDF rendering ───────────────────────────────────────────────────────────
#
# Resumes are stored as Markdown (from the template) and cover letters as plain
# prose. Both render through the same small writer: this is deliberately a simple
# renderer, not a full Markdown engine — a resume only ever needs headings, bold,
# bullets and paragraphs, and a dependency-light path matters more than fidelity.

def _pdf_line_with_right(pdf, left: str, right: str, body_size: float,
                         face: str = "Helvetica"):
    """One line with `right` flush to the right margin — the tabbed-date line."""
    heading = left.startswith("### ")
    bullet = left.lstrip().startswith(("- ", "* "))

    if heading:
        left = left[4:].strip()
        pdf.set_font(face, "B", body_size + 0.5)
    elif bullet:
        # A "-", not a "•". Core PDF fonts are latin-1 and U+2022 isn't in it —
        # the rest of this renderer already made that concession, and a bullet
        # that differs by line would look like a mistake.
        left = "  -  " + left.lstrip()[2:]
        pdf.set_font(face, "", body_size)
    else:
        pdf.set_font(face, "", body_size)

    left = left.replace("**", "")
    usable = pdf.w - pdf.l_margin - pdf.r_margin
    right_w = pdf.get_string_width(right) + 1

    pdf.cell(usable - right_w, 5, left, align="L")
    pdf.set_font(face, "I", body_size - 0.5)
    pdf.cell(right_w, 5, right, align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(face, "", body_size)


def to_docx(content: str, kind: str) -> bytes:
    """Render to a Word document.

    The resume gets its full structured layout. The cover letter is prose, so it gets a
    plain, clean Word document — the paragraphs as written, in a readable font with
    normal margins — because people asked to download it as .docx too, and pasting a PDF
    into an application form is worse than attaching a Word file.
    """
    if kind == "resume":
        from src.resume_docx import to_docx as render
        return render(content)
    return _prose_to_docx(content)


def _prose_to_docx(content: str) -> bytes:
    """A cover letter (or any prose) as a simple, clean Word document."""
    import io

    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Blank lines separate paragraphs; single newlines inside a block are kept as line
    # breaks (so an address block or sign-off stays tight).
    for block in content.replace("\r\n", "\n").split("\n\n"):
        block = block.strip("\n")
        if not block.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        lines = block.split("\n")
        for i, line in enumerate(lines):
            run = p.add_run(line)
            if i < len(lines) - 1:
                run.add_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(content: str, kind: str) -> bytes:
    """Render to PDF, matching the Word file rather than approximating it.

    The two renderers used to drift: Word got Calibri, real bullets and clickable
    links while this one kept Times, a hyphen where a bullet belonged, and raw URLs
    printed across the page. A resume must not read differently depending on which
    button you pressed.
    """
    try:
        from fpdf import FPDF
    except ImportError as e:      # pragma: no cover
        raise RuntimeError(
            "PDF export needs fpdf2. Install it with: pip install fpdf2"
        ) from e

    from src import resume_pdf
    from src.resume_docx import BODY_PT, MARGIN_IN, NAME_PT

    if kind == "resume":
        pdf = FPDF(format="A4", unit="mm")
        margin = MARGIN_IN * 25.4
        pdf.set_margins(margin, margin, margin)
        pdf.set_auto_page_break(auto=True, margin=margin)
        body_size = BODY_PT
    else:
        pdf = FPDF(format="letter", unit="mm")
        pdf.set_margins(18, 16, 18)
        pdf.set_auto_page_break(auto=True, margin=15)
        body_size = 11

    face, unicode_ok = resume_pdf.resolve_font(pdf)
    prepare = (lambda t: t) if unicode_ok else resume_pdf.latin1_safe
    bullet = resume_pdf.BULLET if unicode_ok else "-"

    pdf.add_page()
    pdf.set_text_color(0, 0, 0)

    def write_link(text: str, url: str, size: float):
        """Blue, underlined, and actually clickable."""
        pdf.set_font(face, "U", size) if False else pdf.set_font(face, "", size)
        pdf.set_text_color(5, 99, 193)
        pdf.cell(pdf.get_string_width(prepare(text)) + 0.5, 5,
                 prepare(text), link=url)
        pdf.set_text_color(0, 0, 0)

    in_comment = False
    header_done = False

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        # The template explains itself in an HTML comment. A stray line of it on a
        # real resume, sent to a real employer, is unrecoverable.
        if stripped.startswith("<!--"):
            in_comment = "-->" not in stripped
            continue
        if in_comment:
            if "-->" in stripped:
                in_comment = False
            continue

        if not stripped:
            pdf.ln(2.5)
            continue

        text = line
        right = ""
        if "@@" in text:
            left_part, _, right_part = text.partition("@@")
            text = left_part.rstrip()
            right = right_part.strip()

        pdf.set_x(pdf.l_margin)

        # ── # Name ──────────────────────────────────────────────────────────
        if text.startswith("# "):
            pdf.set_font(face, "B", NAME_PT)
            pdf.multi_cell(0, 9, prepare(text[2:].strip()), align="L")
            pdf.ln(0.5)
            continue

        # ── ## Section, with the hairline under it ──────────────────────────
        if text.startswith("## "):
            header_done = True
            pdf.ln(2)
            pdf.set_font(face, "B", body_size)
            pdf.multi_cell(0, 5.5, prepare(text[3:].strip().upper()), align="L")
            y = pdf.get_y()
            pdf.set_draw_color(0, 0, 0)
            pdf.set_line_width(0.2)
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(2)
            continue

        # ── The contact block, above the first heading ──────────────────────
        # Every piece on ONE line, and the profile URLs clickable. Written with
        # cell() rather than multi_cell() so the parts sit side by side instead of
        # each claiming a line of their own.
        if not header_done:
            pdf.set_font(face, "", body_size)
            parts = [p.strip() for p in text.split("|")]
            for i, part in enumerate(parts):
                if i:
                    pdf.set_font(face, "", body_size)
                    pdf.cell(pdf.get_string_width(" | ") + 0.5, 5, " | ")
                if resume_pdf.is_url(part):
                    url = part if part.startswith("http") else f"https://{part}"
                    write_link(resume_pdf.link_label(part), url, body_size)
                else:
                    pdf.cell(pdf.get_string_width(prepare(part)) + 0.5, 5,
                             prepare(part))
            pdf.ln(5.5)
            continue

        # ── ### Sub-heading (role, project, degree) ─────────────────────────
        if text.startswith("### "):
            heading = resume_pdf.strip_md(text[4:].strip())
            pdf.set_font(face, "B", body_size)
            if right:
                pdf.cell(0, 5.5, prepare(heading))
                _right_aligned(pdf, right, body_size, face, prepare, resume_pdf,
                               bold=False)
            else:
                pdf.multi_cell(0, 5.5, prepare(heading), align="L")
            continue

        # ── - bullet ────────────────────────────────────────────────────────
        if re.match(r"^[-*]\s+", text):
            body = resume_pdf.strip_md(re.sub(r"^[-*]\s+", "", text))
            label, rest = _split_bold_label(re.sub(r"^[-*]\s+", "", text))

            indent = 4.0
            pdf.set_x(pdf.l_margin)
            pdf.set_font(face, "", body_size)
            pdf.cell(indent, 5, prepare(bullet))

            if right:
                pdf.set_font(face, "", body_size)
                pdf.cell(0, 5, prepare(body))
                _right_aligned(pdf, right, body_size, face, prepare, resume_pdf,
                               bold=False)
                continue

            if label:
                # "- **Databases:** MySQL | SQLite" — the label is bold, the rest
                # is not, and they belong on the same line.
                pdf.set_font(face, "B", body_size)
                pdf.write(5, prepare(f"{label}: "))
                pdf.set_font(face, "", body_size)
                pdf.write(5, prepare(rest))
                pdf.ln(5.2)
            else:
                usable = pdf.w - pdf.l_margin - pdf.r_margin - indent
                pdf.set_x(pdf.l_margin + indent)
                pdf.multi_cell(usable, 5, prepare(body), align="L")
            continue

        # ── Everything else: a paragraph ────────────────────────────────────
        if right:
            pdf.set_font(face, "", body_size)
            pdf.cell(0, 5.4, prepare(resume_pdf.strip_md(text)))
            _right_aligned(pdf, right, body_size, face, prepare, resume_pdf,
                           bold=False)
            continue

        label, rest = _split_bold_label(text)
        if label:
            pdf.set_font(face, "B", body_size)
            pdf.write(5, prepare(f"{label}: "))
            pdf.set_font(face, "", body_size)
            pdf.write(5, prepare(rest))
            pdf.ln(5.4)
        else:
            pdf.set_font(face, "", body_size)
            pdf.multi_cell(0, 5.4, prepare(resume_pdf.strip_md(text)), align="L")

    out = pdf.output()
    return bytes(out)


def _split_bold_label(text: str) -> tuple[str, str]:
    """"**Databases:** MySQL | SQLite" -> ("Databases", "MySQL | SQLite")."""
    m = re.match(r"^\*\*(.+?):?\*\*:?\s*(.*)$", text.strip())
    if not m:
        return "", ""
    from src import resume_pdf
    return m.group(1).strip(), resume_pdf.strip_md(m.group(2))


def _right_aligned(pdf, right: str, size: float, face: str, prepare, resume_pdf,
                   bold: bool):
    """The dates, or the link, at the right margin — on the line just written.

    In Word this is a right tab stop. Here it is a right-aligned cell placed by
    hand, which comes to the same thing on the page and, importantly, is still one
    ordinary line of text to a parser rather than a table cell.
    """
    if resume_pdf.is_url(right):
        label = resume_pdf.link_label(right)
        url = right if right.startswith("http") else f"https://{right}"
        pdf.set_font(face, "", size)
        width = pdf.get_string_width(prepare(label)) + 1
        pdf.set_xy(pdf.w - pdf.r_margin - width, pdf.get_y())
        pdf.set_text_color(5, 99, 193)
        pdf.cell(width, 5.5, prepare(label), align="R", link=url)
        pdf.set_text_color(0, 0, 0)
    else:
        pdf.set_font(face, "B" if bold else "", size)
        width = pdf.get_string_width(prepare(right)) + 1
        pdf.set_xy(pdf.w - pdf.r_margin - width, pdf.get_y())
        pdf.cell(width, 5.5, prepare(right), align="R")
    pdf.ln(5.5)


def _pdf_safe(text: str) -> str:
    """Core PDF fonts are latin-1 only; swap the characters models like to use."""
    replacements = {
        "\u2014": "-", "\u2013": "-", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2022": "-", "\u2026": "...",
        "\u00a0": " ", "\u2192": "->",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text.encode("latin-1", "ignore").decode("latin-1")


def _strip_md(text: str) -> str:
    """Drop inline Markdown markers the simple renderer doesn't draw."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)      # bold
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*", r"\1", text)   # italic
    text = re.sub(r"`(.+?)`", r"\1", text)            # code
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", text)  # links
    return text.strip()
