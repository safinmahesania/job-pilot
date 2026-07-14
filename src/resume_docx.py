"""Render the tailored resume to Word, in the shape of your own template.

The styling here is not invented — it is lifted from Resume_Template.docx:
Times New Roman, A4 with half-inch margins, a 20pt small-caps name at the left, an
11pt justified body, and section headings in bold small caps with a black hairline
under them. If you change the template, change these constants to match.

The one deliberate departure is tables. Your template puts dates on the right with
a two-column table, as almost every Word resume does. It is a quiet liability:
Workday and Taleo — which stand between you and a large share of Canadian
employers — are known for scrambling tables, because to a parser a table cell is
not a line of text. Your job title lands in one field and its dates in another.

A right-aligned tab stop puts the date in exactly the same place, and to a parser
the line is just text. Same page, same look, no table.
"""
import re

# ── The template's own styling, measured from Resume_Template.docx ──────────
FONT = "Calibri"

PAGE_WIDTH_IN = 8.27          # A4
PAGE_HEIGHT_IN = 11.69
MARGIN_IN = 0.5

NAME_PT = 20                  # bold, small caps, left
CONTACT_PT = 11
HEADING_PT = 11               # bold, small caps, black rule underneath
BODY_PT = 11                  # justified

HEADING_SPACE_BEFORE_PT = 12
HEADING_SPACE_AFTER_PT = 6

TEXT_WIDTH_IN = PAGE_WIDTH_IN - (2 * MARGIN_IN)

# "left @@ right" — the right half is pushed to the right margin.
RIGHT = "@@"


def _split_right(line: str) -> tuple[str, str]:
    if RIGHT not in line:
        return line.strip(), ""
    left, _, right = line.partition(RIGHT)
    return left.strip(), right.strip()


# A bare URL on a resume is a wall of characters nobody reads and nobody types.
# What a reader wants is a word they can click. So a URL becomes a link labelled by
# what it points AT — the repo, the certificate, the profile — and the label is
# chosen from the URL itself rather than asked of the model, which would get it
# wrong occasionally and unfixably.
def link_label(url: str) -> str:
    text = (url or "").strip()
    lowered = text.lower()

    # The header pair reads better bare — "linkedin.com/in/safinmahesania" tells a
    # reader who you are; "LinkedIn URL" tells them nothing they didn't know.
    for host in ("linkedin.com/in/", "github.com/"):
        if host in lowered:
            path = lowered.split(host, 1)[1].strip("/")
            if path and "/" not in path:            # a profile, not a repo
                return f"{host.rstrip('/')}/{path}"

    if "github.com" in lowered:
        return "GitHub URL"
    if any(k in lowered for k in ("credly", "learn.microsoft", "certificate",
                                  "credential", "drive.google")):
        return "Certificate URL"
    if "linkedin.com" in lowered:
        return "LinkedIn URL"
    return text


def _is_url(text: str) -> bool:
    lowered = (text or "").strip().lower()
    return lowered.startswith(("http://", "https://", "www.")) or (
        "." in lowered and "/" in lowered and " " not in lowered
    )


def _add_hyperlink(paragraph, url: str, label: str, size=BODY_PT, italic=False):
    """A real Word hyperlink. python-docx has no API for this, so it is assembled
    from the relationship up."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    node = OxmlElement("w:hyperlink")
    node.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    props.append(fonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    props.append(sz)

    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "0563C1")           # Word's own hyperlink blue
    props.append(colour)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)

    if italic:
        props.append(OxmlElement("w:i"))

    run.append(props)

    text = OxmlElement("w:t")
    text.text = label
    run.append(text)

    node.append(run)
    paragraph._p.append(node)


def _style(run, size=BODY_PT, bold=False, italic=False, small_caps=False):
    from docx.shared import Pt

    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if small_caps:
        run.font.small_caps = True
    return run


def _add_runs(paragraph, text: str, size=BODY_PT, bold=False):
    """Write text, honouring **bold** spans."""
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        _style(paragraph.add_run(chunk), size=size, bold=bold or (i % 2 == 1))


def _paragraph(doc, space_after=0, space_before=0, justify=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _right_tab(paragraph):
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches

    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(TEXT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT
    )


def _rule(paragraph, color="000000"):
    """The hairline under a section heading — a paragraph border, as in your
    template. Not a one-cell table, which is the usual trick and would put back
    exactly the structure we removed."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def to_docx(markdown: str) -> bytes:
    import io

    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as e:
        raise RuntimeError(
            "Word export needs python-docx. Install it with: pip install python-docx"
        ) from e

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    for side in ("top", "bottom", "left", "right"):
        setattr(section, f"{side}_margin", Inches(MARGIN_IN))

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    normal.paragraph_format.space_after = Pt(0)

    # The template explains itself in an HTML comment. That is for you.
    body = re.sub(r"<!--.*?-->", "", markdown, flags=re.S)

    header_done = False

    for raw in body.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue

        # ── Name: 20pt, bold, small caps, left ──────────────────────────────
        if line.startswith("# "):
            p = _paragraph(doc, space_after=2)
            _style(p.add_run(line[2:].strip()),
                   size=NAME_PT, bold=True, small_caps=True)
            continue

        # ── Section heading: 11pt, bold, small caps, black rule ─────────────
        if line.startswith("## "):
            header_done = True
            p = _paragraph(doc,
                           space_before=HEADING_SPACE_BEFORE_PT,
                           space_after=HEADING_SPACE_AFTER_PT)
            _style(p.add_run(line[3:].strip()),
                   size=HEADING_PT, bold=True, small_caps=True)
            _rule(p)
            continue

        # ── Entry heading: "Role @@ dates" ──────────────────────────────────
        if line.startswith("### "):
            left, right = _split_right(line[4:])
            p = _paragraph(doc, space_before=4, space_after=0)
            _right_tab(p)
            _add_runs(p, left, size=BODY_PT, bold=True)
            if right:
                p.add_run("\t")
                if _is_url(right):
                    _add_hyperlink(p, right, link_label(right))
                else:
                    _style(p.add_run(right), size=BODY_PT)
            continue

        # ── Bullet ──────────────────────────────────────────────────────────
        if line.lstrip().startswith(("- ", "* ")):
            text = line.lstrip()[2:]
            left, right = _split_right(text)

            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.25)

            if right:
                _right_tab(p)
                _add_runs(p, left)
                p.add_run("\t")
                if _is_url(right):
                    _add_hyperlink(p, right, link_label(right))
                else:
                    _style(p.add_run(right), size=BODY_PT)
            else:
                _add_runs(p, left)
            continue

        # ── Contact block, above the first heading ──────────────────────────
        if not header_done:
            # LEFT, not justified. Justifying a short line stretches the spaces
            # inside it until "+1 437 661 5569" reads as four separate numbers,
            # which is what a phone number must never do.
            p = _paragraph(doc, space_after=0, justify=False)
            parts = [part.strip() for part in line.split("|")]
            for i, part in enumerate(parts):
                if i:
                    _style(p.add_run(" | "), size=CONTACT_PT)
                if _is_url(part):
                    url = part if part.startswith("http") else f"https://{part}"
                    _add_hyperlink(p, url, link_label(part), size=CONTACT_PT)
                else:
                    _add_runs(p, part, size=CONTACT_PT)
            continue

        # ── Body prose (summary, company, university, volunteer copy) ───────
        left, right = _split_right(line)
        if right:
            p = _paragraph(doc, space_after=0)
            _right_tab(p)
            _add_runs(p, left)
            p.add_run("\t")
            if _is_url(right):
                _add_hyperlink(p, right, link_label(right))
            else:
                _style(p.add_run(right), size=BODY_PT)
        else:
            # Justified, as in your template.
            p = _paragraph(doc, space_after=0, justify=True)
            _add_runs(p, left)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
